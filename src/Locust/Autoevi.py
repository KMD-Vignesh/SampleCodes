import os
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# Paths
testcase_excel = "test_config.csv"  # Path to the Test Case Excel
base_dir = r"data\base_dir\MLEGALAlertData"  # Base directory containing test folders
output_docx = "test_case_output.docx"  # Path to save the Word document
file_patterns = ["dxQuote", "dxTrade"]  # Replace with your file patterns to filter files

# Read the test case Excel file
testcase_df = pd.read_csv(testcase_excel)  # Assuming the test_config.csv is a CSV file

# Create a Word document
document = Document()

# Loop through each test case in the Excel file
for _, row in testcase_df.iterrows():
    testcase_id = row["TestID"]
    description = row["Description"]
    folder_path = os.path.join(base_dir, testcase_id)
    
    # Add Test Case ID and Description to Word document
    document.add_heading(f"Test Case: {testcase_id}", level=1)
    document.add_paragraph(f"Description: {description}")
    
    # Check if the folder exists
    if os.path.isdir(folder_path):
        # Loop through files in the folder
        for file_name in os.listdir(folder_path):
            # Check if the file matches any of the patterns
            if any(pattern in file_name for pattern in file_patterns):
                file_path = os.path.join(folder_path, file_name)
                
                # Read the CSV file
                if file_name.endswith(".csv"):
                    try:
                        df = pd.read_csv(file_path)
                        
                        # Create a screenshot of the CSV file
                        plt.figure(figsize=(8, len(df) * 0.4))  # Adjust the figure size
                        plt.axis("off")
                        table = plt.table(cellText=df.values, colLabels=df.columns, loc="center", cellLoc="center")
                        table.auto_set_font_size(False)
                        table.set_fontsize(8)
                        table.auto_set_column_width(col=list(range(len(df.columns))))
                        
                        # Save the screenshot as a temporary image
                        screenshot_path = f"{file_path}.png"
                        plt.savefig(screenshot_path, bbox_inches="tight")
                        plt.close()
                        
                        # Add File Name and Screenshot to the Word document
                        document.add_heading(file_name, level=2)
                        document.add_picture(screenshot_path, width=Inches(6))
                        
                        # Remove the temporary image file
                        os.remove(screenshot_path)
                    except Exception as e:
                        document.add_paragraph(f"Could not process file '{file_name}'. Error: {e}", style="Italic")
    else:
        # Add a note if the folder does not exist
        document.add_paragraph(f"Folder '{testcase_id}' not found.", style="Italic")

# Save the Word document
document.save(output_docx)
print(f"Word document saved as {output_docx}")
